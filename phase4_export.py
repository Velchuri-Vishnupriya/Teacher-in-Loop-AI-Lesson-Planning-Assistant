import io
import re

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
)

from reportlab.lib.styles import (
    getSampleStyleSheet,
)


# =====================================================
# MARKDOWN CLEANER
# =====================================================

def clean_markdown(text):

    # Remove bold markdown
    text = re.sub(
        r"\*\*(.*?)\*\*",
        r"\1",
        text,
    )

    # Remove italic markdown
    text = re.sub(
        r"\*(.*?)\*",
        r"\1",
        text,
    )

    return text


# =====================================================
# DOCX EXPORT
# =====================================================

def create_docx(lesson_plan):

    document = Document()

    document.add_heading(
        "Teacher-in-the-Loop Lesson Plan",
        level=1,
    )

    lines = lesson_plan.split("\n")

    for line in lines:

        line = line.strip()

        if not line:
            continue

        # -------------------------
        # H2 Heading (##)
        # -------------------------

        if line.startswith("##"):

            heading = (
                line.replace(
                    "##",
                    "",
                )
                .strip()
            )

            document.add_heading(
                heading,
                level=2,
            )

        # -------------------------
        # H1 Heading (#)
        # -------------------------

        elif line.startswith("#"):

            heading = (
                line.replace(
                    "#",
                    "",
                )
                .strip()
            )

            document.add_heading(
                heading,
                level=1,
            )

        # -------------------------
        # Normal Text
        # -------------------------

        else:

            line = clean_markdown(line)

            document.add_paragraph(
                line
            )

    buffer = io.BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer


# =====================================================
# PDF EXPORT
# =====================================================

def create_pdf(lesson_plan):

    buffer = io.BytesIO()

    pdf = SimpleDocTemplate(
        buffer
    )

    styles = getSampleStyleSheet()

    content = []

    content.append(
        Paragraph(
            "Teacher-in-the-Loop Lesson Plan",
            styles["Title"],
        )
    )

    content.append(
        Spacer(
            1,
            12,
        )
    )

    lines = lesson_plan.split("\n")

    for line in lines:

        line = line.strip()

        if not line:
            continue

        # -------------------------
        # H2 Heading (##)
        # -------------------------

        if line.startswith("##"):

            heading = (
                line.replace(
                    "##",
                    "",
                )
                .strip()
            )

            content.append(
                Paragraph(
                    heading,
                    styles["Heading2"],
                )
            )

        # -------------------------
        # H1 Heading (#)
        # -------------------------

        elif line.startswith("#"):

            heading = (
                line.replace(
                    "#",
                    "",
                )
                .strip()
            )

            content.append(
                Paragraph(
                    heading,
                    styles["Heading1"],
                )
            )

        # -------------------------
        # Normal Text
        # -------------------------

        else:

            line = clean_markdown(line)

            content.append(
                Paragraph(
                    line,
                    styles["BodyText"],
                )
            )

        content.append(
            Spacer(
                1,
                6,
            )
        )

    pdf.build(content)

    buffer.seek(0)

    return buffer